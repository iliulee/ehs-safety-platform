# -*- coding: utf-8 -*-
"""
A3分包管理模板批量添加变量标记工具
自动将固定的文字替换为 {{变量名}} 格式
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# ==================== 配置区 ====================

# 原文件目录
SOURCE_DIR = r"f:\20260517app\一分公司内业资料\A3分包管理\A-3-2分包单位资料"

# 输出目录
OUTPUT_DIR = r"f:\20260517app\一分公司内业资料\A3分包管理\converted_templates"

# 要处理的文件（留空则处理全部docx）
TARGET_FILES = [
    "A-3-2-1总包对分包交底(自有分包、甲指分包）.docx",
    "A-3-2-2安全生产协议书（自有分包、甲指分包）.docx",
    "A-3-2-2安全生产协议书（甲直分包）.docx",
    "A-3-2-3消防责任书（所有分包都需要，注意称呼修改）.docx",
    "A-3-2-4印章授权书（有项目印章的分包需要）.docx",
    "A-3-2-5风险告知书（甲直分包需要）.docx",
    "A-3-2-7法人授权委托书（所有分包都需要）.docx",
    "A-3-2-8分包单位安全管理体系（所有分包都需要）.docx",
    "A-3-2-9管理人员名单（所有分包都需要）.docx",
]

# ==================== 替换规则 ====================

REPLACEMENTS = [
    # 项目名称
    (r'XX工程', '{{project_name}}'),
    (r'XXXX项目', '{{project_name}}'),
    (r'XXX工程', '{{project_name}}'),
    (r'XXXX工程', '{{project_name}}'),
    
    # 施工单位
    (r'云南机场建设发展有限公司', '{{company}}'),
    (r'北京佳和建设工程有限公司', '{{company}}'),
    (r'北京佳和建设', '{{company}}'),
    
    # 项目经理
    (r'项目经理：\S+', '项目经理：{{project_manager}}'),
    (r'项目经理：\S+\s', '项目经理：{{project_manager}} '),
    
    # 技术负责人
    (r'技术负责人：\S+', '技术负责人：{{tech_manager}}'),
    
    # 安全负责人
    (r'安全负责人：\S+', '安全负责人：{{safety_manager}}'),
    
    # 项目部
    (r'XX工程项目部', '{{project_name}}项目部'),
    (r'XXXX项目部', '{{project_name}}项目部'),
    
    # 分包单位
    (r'XX建筑公司', '{{subcontractor}}'),
    (r'XX建筑工程有限公司', '{{subcontractor}}'),
    (r'XX劳务公司', '{{subcontractor}}'),
    (r'XX科技有限公司', '{{subcontractor}}'),
    
    # 分包负责人
    (r'分包单位负责人：\S+', '分包单位负责人：{{subcontractor_manager}}'),
    
    # 日期处理
    (r'(\d{4})年', '{{year}}年'),
    (r'(\d{1,2})月(\d{1,2})日', '{{month}}月{{day}}日'),
    (r'月(\d{1,2})日', '月{{day}}日'),
    
    # 地址
    (r'工程地点：\S+', '工程地点：{{address}}'),
    
    # 甲方乙方
    (r'甲方：\S+公司', '甲方：{{company}}'),
    (r'乙方：\S+公司', '乙方：{{subcontractor}}'),
]

# ==================== 函数区 ====================

def process_paragraph(para):
    """处理单个段落"""
    original = para.text
    modified = original
    
    for pattern, replacement in REPLACEMENTS:
        modified = re.sub(pattern, replacement, modified)
    
    if modified != original:
        # 替换段落中的文字（保留格式）
        for run in para.runs:
            for pattern, replacement in REPLACEMENTS:
                if re.search(pattern, run.text):
                    run.text = re.sub(pattern, replacement, run.text)

def process_table(table):
    """处理表格"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                process_paragraph(para)

def process_document(input_path, output_path):
    """处理单个文档"""
    try:
        doc = Document(input_path)
        
        # 处理所有段落
        for para in doc.paragraphs:
            process_paragraph(para)
        
        # 处理所有表格
        for table in doc.tables:
            process_table(table)
        
        # 保存到输出目录
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return True, "成功"
    except Exception as e:
        return False, str(e)

def main():
    """主函数"""
    print("=" * 60)
    print("A3分包管理模板批量添加变量标记工具")
    print("=" * 60)
    print()
    
    # 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 如果没有指定文件，获取所有docx文件
    if not TARGET_FILES:
        files = [f for f in os.listdir(SOURCE_DIR) if f.endswith('.docx')]
    else:
        files = TARGET_FILES
    
    # 统计
    total = len(files)
    success = 0
    failed = 0
    failed_files = []
    
    print(f"源目录: {SOURCE_DIR}")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"待处理文件数: {total}")
    print()
    print("-" * 60)
    
    # 处理每个文件
    for i, filename in enumerate(files, 1):
        input_path = os.path.join(SOURCE_DIR, filename)
        output_path = os.path.join(OUTPUT_DIR, filename)
        
        if not os.path.exists(input_path):
            print(f"[{i}/{total}] 跳过（文件不存在）: {filename}")
            continue
        
        print(f"[{i}/{total}] 处理中: {filename}")
        
        ok, msg = process_document(input_path, output_path)
        
        if ok:
            success += 1
            print(f"    ✓ 转换成功")
        else:
            failed += 1
            failed_files.append((filename, msg))
            print(f"    ✗ 失败: {msg}")
    
    # 输出统计
    print()
    print("-" * 60)
    print("处理完成！")
    print(f"总计: {total} 个文件")
    print(f"成功: {success} 个")
    print(f"失败: {failed} 个")
    
    if failed_files:
        print()
        print("失败文件列表:")
        for filename, error in failed_files:
            print(f"  - {filename}: {error}")
    
    print()
    print(f"输出目录: {OUTPUT_DIR}")
    print()
    
    input("按回车键退出...")

if __name__ == "__main__":
    main()
